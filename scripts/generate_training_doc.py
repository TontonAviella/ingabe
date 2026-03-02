#!/usr/bin/env python3
"""Generate a professional Word document from the Ingabe Training Manual markdown."""

import re
from pathlib import Path

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor

BRAND_GREEN = RGBColor(0xA3, 0xE6, 0x35)  # Ingabe green from the app
BRAND_DARK = RGBColor(0x1E, 0x1E, 0x2E)
HEADING_COLOR = RGBColor(0x2D, 0x5A, 0x27)  # Dark green for headings
SUBHEADING_COLOR = RGBColor(0x3A, 0x7D, 0x44)
TEXT_COLOR = RGBColor(0x33, 0x33, 0x33)
LIGHT_GREEN_BG = "E8F5E9"
TABLE_HEADER_BG = "2D5A27"
TABLE_ALT_BG = "F1F8E9"


def set_cell_shading(cell, color_hex: str):
    """Set background color on a table cell."""
    shading = cell._element.get_or_add_tcPr()
    shading_elem = shading.makeelement(
        qn("w:shd"),
        {
            qn("w:fill"): color_hex,
            qn("w:val"): "clear",
        },
    )
    shading.append(shading_elem)


def style_table(table, has_header=True):
    """Apply consistent styling to a table."""
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Style header row
    if has_header and table.rows:
        for cell in table.rows[0].cells:
            set_cell_shading(cell, TABLE_HEADER_BG)
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
                    run.font.bold = True
                    run.font.size = Pt(9)

    # Style alternating rows
    for i, row in enumerate(table.rows):
        if i == 0 and has_header:
            continue
        if i % 2 == 0:
            for cell in row.cells:
                set_cell_shading(cell, TABLE_ALT_BG)

    # Set font size for all cells
    for row in table.rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                paragraph.paragraph_format.space_before = Pt(2)
                paragraph.paragraph_format.space_after = Pt(2)
                for run in paragraph.runs:
                    if run.font.size is None:
                        run.font.size = Pt(9)


def add_styled_paragraph(doc, text, style="Normal", bold=False, color=None, size=None):
    """Add a paragraph with optional styling."""
    p = doc.add_paragraph(style=style)
    run = p.add_run(text)
    if bold:
        run.font.bold = True
    if color:
        run.font.color.rgb = color
    if size:
        run.font.size = size
    return p


def parse_table(lines: list[str]) -> tuple[list[str], list[list[str]]]:
    """Parse a markdown table into headers and rows."""
    headers = []
    rows = []
    for i, line in enumerate(lines):
        cells = [c.strip() for c in line.strip().strip("|").split("|")]
        if i == 0:
            headers = cells
        elif i == 1:
            continue  # separator line
        else:
            rows.append(cells)
    return headers, rows


def add_table_to_doc(doc, headers: list[str], rows: list[list[str]]):
    """Add a formatted table to the document."""
    ncols = len(headers)
    table = doc.add_table(rows=1 + len(rows), cols=ncols)
    table.style = "Table Grid"

    # Fill header
    for j, header in enumerate(headers):
        cell = table.rows[0].cells[j]
        cell.text = header
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Fill data rows
    for i, row_data in enumerate(rows):
        for j in range(min(len(row_data), ncols)):
            cell = table.rows[i + 1].cells[j]
            text = row_data[j]
            # Handle bold markdown
            if text.startswith("**") and text.endswith("**"):
                p = cell.paragraphs[0]
                run = p.add_run(text.strip("*"))
                run.font.bold = True
            else:
                cell.text = text.strip("*")

    style_table(table)
    doc.add_paragraph()  # spacing after table


def build_document():
    """Build the complete Word document."""
    doc = Document()

    # ---- Page setup ----
    section = doc.sections[0]
    section.page_width = Cm(21)  # A4
    section.page_height = Cm(29.7)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)
    section.top_margin = Cm(2)
    section.bottom_margin = Cm(2)

    # ---- Modify default styles ----
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(10.5)
    style.font.color.rgb = TEXT_COLOR
    style.paragraph_format.space_after = Pt(6)
    style.paragraph_format.line_spacing = 1.15

    for level in range(1, 4):
        h_style = doc.styles[f"Heading {level}"]
        h_style.font.name = "Calibri"
        if level == 1:
            h_style.font.size = Pt(22)
            h_style.font.color.rgb = HEADING_COLOR
            h_style.font.bold = True
            h_style.paragraph_format.space_before = Pt(24)
            h_style.paragraph_format.space_after = Pt(12)
        elif level == 2:
            h_style.font.size = Pt(16)
            h_style.font.color.rgb = HEADING_COLOR
            h_style.font.bold = True
            h_style.paragraph_format.space_before = Pt(18)
            h_style.paragraph_format.space_after = Pt(8)
        elif level == 3:
            h_style.font.size = Pt(13)
            h_style.font.color.rgb = SUBHEADING_COLOR
            h_style.font.bold = True
            h_style.paragraph_format.space_before = Pt(12)
            h_style.paragraph_format.space_after = Pt(6)

    # ========== COVER PAGE ==========
    for _ in range(6):
        doc.add_paragraph()

    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title_p.add_run("INGABE")
    run.font.size = Pt(48)
    run.font.bold = True
    run.font.color.rgb = HEADING_COLOR

    subtitle_p = doc.add_paragraph()
    subtitle_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle_p.add_run("Training Manual")
    run.font.size = Pt(28)
    run.font.color.rgb = SUBHEADING_COLOR

    doc.add_paragraph()

    desc_p = doc.add_paragraph()
    desc_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = desc_p.add_run(
        "A Complete Guide for Farmers, Agronomists,\nNGOs, and Agricultural Professionals"
    )
    run.font.size = Pt(14)
    run.font.color.rgb = TEXT_COLOR

    for _ in range(4):
        doc.add_paragraph()

    footer_p = doc.add_paragraph()
    footer_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = footer_p.add_run("Developed by NozaLabs\nhttps://gis.nozalabs.rw")
    run.font.size = Pt(11)
    run.font.color.rgb = SUBHEADING_COLOR

    doc.add_page_break()

    # ========== PARSE MARKDOWN AND BUILD BODY ==========
    md_path = Path(__file__).parent.parent / "docs" / "TRAINING_MANUAL.md"
    content = md_path.read_text(encoding="utf-8")
    lines = content.split("\n")

    i = 0
    in_table = False
    table_lines: list[str] = []
    in_code_block = False
    code_lines: list[str] = []

    while i < len(lines):
        line = lines[i]

        # Skip the first title and subtitle (we have a cover page)
        if i < 5 and (line.startswith("# Ingabe") or line.startswith("**A Complete Guide")):
            i += 1
            continue

        # Code block handling
        if line.strip().startswith("```"):
            if in_code_block:
                # End code block
                code_text = "\n".join(code_lines)
                p = doc.add_paragraph()
                run = p.add_run(code_text)
                run.font.name = "Consolas"
                run.font.size = Pt(9)
                run.font.color.rgb = RGBColor(0x2E, 0x7D, 0x32)
                p.paragraph_format.left_indent = Cm(1)
                in_code_block = False
                code_lines = []
            else:
                in_code_block = True
                code_lines = []
            i += 1
            continue

        if in_code_block:
            code_lines.append(line)
            i += 1
            continue

        # Table handling
        if "|" in line and line.strip().startswith("|"):
            if not in_table:
                in_table = True
                table_lines = []
            table_lines.append(line)
            i += 1
            continue
        elif in_table:
            # End of table
            headers, rows = parse_table(table_lines)
            if headers and rows:
                add_table_to_doc(doc, headers, rows)
            in_table = False
            table_lines = []
            # Don't increment i, process current line

        # Horizontal rule / section break
        if line.strip() == "---":
            i += 1
            continue

        # Headings
        if line.startswith("## "):
            # Section heading (H1 in the doc, since we skip the markdown H1)
            heading_text = line[3:].strip()
            # Remove markdown link syntax
            heading_text = re.sub(r"\[([^\]]+)\]\([^)]+\)", r"\1", heading_text)
            # Remove numbering prefix like "1. "
            heading_text = re.sub(r"^\d+\.\s+", "", heading_text)
            doc.add_heading(heading_text, level=1)
            i += 1
            continue

        if line.startswith("### "):
            heading_text = line[4:].strip()
            heading_text = re.sub(r"\[([^\]]+)\]\([^)]+\)", r"\1", heading_text)
            doc.add_heading(heading_text, level=2)
            i += 1
            continue

        if line.startswith("#### "):
            heading_text = line[5:].strip()
            doc.add_heading(heading_text, level=3)
            i += 1
            continue

        # Empty line
        if not line.strip():
            i += 1
            continue

        # Table of Contents lines (skip — we generate our own)
        if re.match(r"^\d+\.\s+\[", line):
            i += 1
            continue

        # Bullet points
        if line.strip().startswith("- "):
            text = line.strip()[2:]
            p = doc.add_paragraph(style="List Bullet")
            # Handle bold within bullets
            parts = re.split(r"(\*\*[^*]+\*\*)", text)
            for part in parts:
                if part.startswith("**") and part.endswith("**"):
                    run = p.add_run(part[2:-2])
                    run.font.bold = True
                else:
                    # Handle inline code
                    code_parts = re.split(r"(`[^`]+`)", part)
                    for cp in code_parts:
                        if cp.startswith("`") and cp.endswith("`"):
                            run = p.add_run(cp[1:-1])
                            run.font.name = "Consolas"
                            run.font.size = Pt(9)
                            run.font.color.rgb = RGBColor(0x2E, 0x7D, 0x32)
                        else:
                            p.add_run(cp)
            i += 1
            continue

        # Numbered list
        match = re.match(r"^(\d+)\.\s+(.*)", line.strip())
        if match:
            text = match.group(2)
            p = doc.add_paragraph(style="List Number")
            # Handle bold and inline code
            parts = re.split(r"(\*\*[^*]+\*\*)", text)
            for part in parts:
                if part.startswith("**") and part.endswith("**"):
                    run = p.add_run(part[2:-2])
                    run.font.bold = True
                else:
                    code_parts = re.split(r"(`[^`]+`)", part)
                    for cp in code_parts:
                        if cp.startswith("`") and cp.endswith("`"):
                            run = p.add_run(cp[1:-1])
                            run.font.name = "Consolas"
                            run.font.size = Pt(9)
                        else:
                            p.add_run(cp)
            i += 1
            continue

        # Regular paragraphs
        text = line.strip()
        if text:
            p = doc.add_paragraph()
            # Handle bold, italic, inline code, and links
            parts = re.split(r"(\*\*[^*]+\*\*|`[^`]+`|\[[^\]]+\]\([^)]+\))", text)
            for part in parts:
                if part.startswith("**") and part.endswith("**"):
                    run = p.add_run(part[2:-2])
                    run.font.bold = True
                elif part.startswith("`") and part.endswith("`"):
                    run = p.add_run(part[1:-1])
                    run.font.name = "Consolas"
                    run.font.size = Pt(9)
                    run.font.color.rgb = RGBColor(0x2E, 0x7D, 0x32)
                elif part.startswith("["):
                    link_match = re.match(r"\[([^\]]+)\]\(([^)]+)\)", part)
                    if link_match:
                        run = p.add_run(link_match.group(1))
                        run.font.color.rgb = RGBColor(0x1B, 0x5E, 0x20)
                        run.font.underline = True
                    else:
                        p.add_run(part)
                else:
                    p.add_run(part)

        i += 1

    # Handle any remaining table
    if in_table and table_lines:
        headers, rows = parse_table(table_lines)
        if headers and rows:
            add_table_to_doc(doc, headers, rows)

    # ========== FOOTER ==========
    doc.add_page_break()
    end_p = doc.add_paragraph()
    end_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for _ in range(8):
        end_p.add_run("\n")
    run = end_p.add_run("Ingabe Training Manual")
    run.font.size = Pt(16)
    run.font.bold = True
    run.font.color.rgb = HEADING_COLOR

    end_p2 = doc.add_paragraph()
    end_p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = end_p2.add_run(
        "Developed by NozaLabs\n"
        "https://gis.nozalabs.rw\n\n"
        "For questions or support, visit https://app.nozalabs.rw"
    )
    run.font.size = Pt(11)
    run.font.color.rgb = TEXT_COLOR

    # ========== SAVE ==========
    out_path = Path(__file__).parent.parent / "docs" / "Ingabe_Training_Manual.docx"
    doc.save(str(out_path))
    print(f"Document saved to: {out_path}")
    print(f"File size: {out_path.stat().st_size / 1024:.0f} KB")


if __name__ == "__main__":
    build_document()
